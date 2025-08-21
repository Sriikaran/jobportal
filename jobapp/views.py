from django.shortcuts import render, redirect
from django.http import JsonResponse
from .models import Login, Employer, JobSeeker, Enquiry
from adminapp.models import News
from employer.models import Jobs
from django.core.exceptions import ObjectDoesNotExist
from django.contrib import messages
from django.urls import reverse
import datetime
from django.db.models import Q
from django.contrib import messages

def test_template(request):
    return render(request, ".html")

def index(request):
    jobs = Jobs.objects.all()
    q = request.GET.get('q')
    if q:
        jobs = jobs.filter(
            Q(post__icontains=q) |
            Q(location__icontains=q) |
            Q(jobdesc__icontains=q) |
            Q(firmname__icontains=q) |
            Q(qualification__icontains=q)
        )
        jobs = list(jobs)  # Force evaluation
        print(f"Search query: {q}, Found jobs: {len(jobs)}, Jobs: {[job.post for job in jobs]}")
    else:
        print(f"No search query, Total jobs: {jobs.count()}")
    context = {'jobs': jobs, 'q': q}
    print("Context:", context)
    return render(request, "index.html", context)

def search_jobs(request):
    jobs = Jobs.objects.all()
    q = request.GET.get('q')
    if q:
        jobs = jobs.filter(
            Q(post__icontains=q) |
            Q(location__icontains=q) |
            Q(jobdesc__icontains=q) |
            Q(firmname__icontains=q) |
            Q(qualification__icontains=q)
        )
    job_list = [{
        'post': job.post,
        'location': job.location,
        'posteddate': job.posteddate,  # Use string directly
        'jobtitle': job.jobtitle,
        'firmname': job.firmname,
        'qualification': job.qualification,
        'jobdesc': job.jobdesc,
        'salarypa': str(job.salarypa),
        'emailaddress': job.emailaddress
    } for job in jobs]
    return JsonResponse({'jobs': job_list})

def aboutus(request):
    return render(request, "aboutus.html")

from django.shortcuts import render
from django.contrib import messages
import datetime

def adminreg(request):
    if request.method == "POST":
        profilepic = request.FILES.get("profilepic")
        fullName = request.POST["fullName"]
        username = request.POST["username"]
        password = request.POST["password"]
        confirmPassword = request.POST["confirmPassword"]
        email = request.POST["email"]
        phone = request.POST["phone"]
        gender = request.POST["gender"]
        dob = request.POST["dob"]
        address = request.POST["address"]
        city = request.POST["city"]
        state = request.POST["state"]
        country = request.POST["country"]
        employeeId = request.POST["employeeId"]
        department = request.POST["department"]
        accessLevel = request.POST["accessLevel"]
        regdate = datetime.datetime.today()

        # Validate password confirmation
        if password != confirmPassword:
            messages.error(request, "Passwords do not match!")
            return render(request, "admin_registration.html", {"msg": "Passwords do not match!"})

        try:
            # Create Administrator record
            admin = Administrator(
                profilepic=profilepic,
                fullName=fullName,
                username=username,
                email=email,
                phone=phone,
                gender=gender,
                dob=dob,
                address=address,
                city=city,
                state=state,
                country=country,
                employeeId=employeeId,
                department=department,
                accessLevel=accessLevel,
                regdate=regdate
            )
            admin.save()

            # Create Login record
            log = Login(
                username=username, 
                password=password, 
                usertype="administrator"
            )
            log.save()

            messages.success(request, "Administrator registration successful! You can now log in.")
            return render(request, "admin_registration.html", {"show_modal": True})

        except Exception as e:
            messages.error(request, f"Registration failed: {str(e)}")
            return render(request, "admin_registration.html", {"msg": f"Registration failed: {str(e)}"})

    return render(request, "administrator.html")


def jobseekerreg(request):
    if request.method == "POST":
        profilepic = request.FILES.get("profilepic")
        name = request.POST["name"]
        gender = request.POST["gender"]
        address = request.POST["address"]
        contactno = request.POST["contactno"]
        emailaddress = request.POST["emailaddress"]
        password = request.POST["password"]
        dob = request.POST["dob"]
        regdate = datetime.datetime.today()

        jobseek = JobSeeker(
            profilepic=profilepic,
            name=name,
            gender=gender,
            address=address,
            contactno=contactno,
            emailaddress=emailaddress,
            dob=dob,
            regdate=regdate
        )
        jobseek.save()

        log = Login(username=emailaddress, password=password, usertype="jobseeker")
        log.save()

        messages.success(request, "Registration successful! You can now log in.")
        return render(request, "jobseeker.html", {"show_modal": True})

    return render(request, "jobseeker.html")

def login(request):
    if request.method == "POST":
        username = request.POST.get("username")
        password = request.POST.get("password")
        try:
            obj = Login.objects.get(username=username)
        except ObjectDoesNotExist:
            return render(request, 'login.html', {"msg": "No user found with this username."})

        if obj.password.strip() != password.strip():
            return render(request, 'login.html', {"msg": "Incorrect password."})

        usertype = obj.usertype
        request.session["username"] = username
        request.session["usertype"] = usertype

        if usertype == 'jobseeker':
            return redirect("jsapp:jshome")
        elif usertype == 'administrator':
            request.session["adminid"] = username
            return redirect("adminapp:adminhome")
        elif usertype == 'employer':
            return redirect("employer:employerhome")
        else:
            return render(request, 'login.html', {"msg": "Invalid user type."})

    return render(request, "login.html")

def employerreg(request):
    if request.method == "POST":
        empprofilepic = request.FILES.get("empprofilepic")
        firmname = request.POST["firmname"]
        firmwork = request.POST["firmwork"]
        firmaddress = request.POST["firmaddress"]
        cpname = request.POST["cpname"]
        cpcontactno = request.POST["cpcontactno"]
        cpemailaddress = request.POST["cpemailaddress"]
        password = request.POST["password"]
        aadharno = request.POST["aadharno"]
        panno = request.POST["panno"]
        gstno = request.POST["gstno"]
        regdate = datetime.datetime.today()

        empreg = Employer(
            empprofilepic=empprofilepic,
            firmname=firmname,
            firmwork=firmwork,
            firmaddress=firmaddress,
            cpname=cpname,
            cpcontactno=cpcontactno,
            cpemailaddress=cpemailaddress,
            aadharno=aadharno,
            panno=panno,
            gstno=gstno,
            regdate=regdate
        )
        empreg.save()

        log = Login(username=cpemailaddress, password=password, usertype="employer")
        log.save()

        messages.success(request, "Employer registration successful! You are now logged in.")
        return redirect("employer:employerhome")  # ✅ Go straight to employer home

    return render(request, "employer.html")


def login(request):

    if request.method == "POST":
        username = request.POST["username"]
        password = request.POST["password"]

        try:
            obj = Login.objects.get(username=username, password=password)
            request.session["username"] = username
            request.session["usertype"] = obj.usertype 
            print("SESSION:", request.session.items())
 

            if obj.usertype == 'jobseeker':
                return redirect("jsapp:jshome")
            elif obj.usertype == 'administrator':
                return redirect("adminapp:adminhome")
            elif obj.usertype == 'employer':
                return redirect("employer:employerhome")
            else:
                return render(request, 'login.html', {"msg": "Invalid user type"})

        except Login.DoesNotExist:
            return render(request, 'login.html', {"msg": "Invalid username or password"})

    return render(request, "login.html")



def contactus(request):
    if request.method=="POST": 
        name=request.POST["name"] 
        gender=request.POST["gender"]
        address=request.POST["address"]
        contactno=request.POST["contactno"]
        emailaddress=request.POST["emailaddress"]
        enquirytext=request.POST["enquirytext"]
        posteddate=datetime.datetime.today()
        enq=Enquiry(name=name, gender=gender, address=address, contactno=contactno, emailaddress=emailaddress, enquirytext=enquirytext, posteddate=posteddate)
        enq.save()
        return render(request, "contactus.html", {"msg": "Enquiry is saved"})
    return render(request, "contactus.html")

def apply(request):
    return render(request, "apply.html")

def services(request):
    return render(request, "services.html")

def blog(request):
    new = News.objects.all()
    return render(request, "blog.html", locals())

def forgot_password(request):
    if request.method == 'POST':
        email = request.POST.get('email')
        try:
            user = Login.objects.get(username=email)
            return render(request, 'reset_password.html', {'user_email': user.username})
        except Login.DoesNotExist:
            messages.error(request, "Email not registered.")
    return render(request, 'forgot_password.html')

from django.shortcuts import render
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_http_methods
import os
import re
import json
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile

# Try importing PDF and DOCX libraries (install with: pip install PyPDF2 python-docx)
try:
    from PyPDF2 import PdfReader
    HAS_PDF_SUPPORT = True
except ImportError:
    HAS_PDF_SUPPORT = False

try:
    import docx
    HAS_DOCX_SUPPORT = True
except ImportError:
    HAS_DOCX_SUPPORT = False

ALLOWED_EXTENSIONS = ['.pdf', '.docx', '.txt']
MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB

class ATSResumeChecker:
    def __init__(self):
        # Keywords that indicate good ATS formatting
        self.required_sections = [
            'experience', 'education', 'skills', 'work experience', 
            'employment', 'qualifications', 'background'
        ]
        
        self.preferred_headings = [
            'professional experience', 'work history', 'employment history',
            'education', 'academic background', 'qualifications',
            'skills', 'technical skills', 'core competencies',
            'summary', 'professional summary', 'objective'
        ]
        
        # Common ATS-unfriendly elements
        self.problematic_elements = [
            'table', 'image', 'graphic', 'chart', 'picture',
            'header', 'footer', 'text box', 'column'
        ]
        
        # Industry keywords for different fields
        self.tech_keywords = [
            'python', 'java', 'javascript', 'react', 'angular', 'sql',
            'aws', 'docker', 'kubernetes', 'agile', 'scrum', 'git'
        ]
        
        self.business_keywords = [
            'management', 'leadership', 'strategy', 'analysis',
            'project management', 'team lead', 'operations'
        ]

    def extract_text_from_file(self, file, file_extension):
        """Extract text content from uploaded file"""
        text_content = ""
        
        try:
            if file_extension == ".pdf" and HAS_PDF_SUPPORT:
                reader = PdfReader(file)
                for page in reader.pages:
                    text_content += page.extract_text() or ""
            
            elif file_extension == ".docx" and HAS_DOCX_SUPPORT:
                doc = docx.Document(file)
                for para in doc.paragraphs:
                    text_content += para.text + "\n"
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text_content += cell.text + "\n"
            
            elif file_extension == ".txt":
                text_content = file.read().decode("utf-8")
                
        except Exception as e:
            raise Exception(f"Error extracting text from file: {str(e)}")
        
        return text_content

    def analyze_resume(self, text_content, filename):
        """Perform comprehensive ATS analysis"""
        analysis_result = {
            "file_name": filename,
            "ats_score": 100,
            "issues": [],
            "recommendations": [],
            "word_count": 0,
            "sections_found": [],
            "keywords_found": []
        }
        
        if not text_content.strip():
            analysis_result["ats_score"] = 0
            analysis_result["issues"].append("No readable text content found in the file")
            return analysis_result
        
        text_lower = text_content.lower()
        words = text_content.split()
        analysis_result["word_count"] = len(words)
        
        # Check for required sections
        sections_found = 0
        for section in self.required_sections:
            if section in text_lower:
                sections_found += 1
                analysis_result["sections_found"].append(section.title())
        
        if sections_found < 2:
            analysis_result["ats_score"] -= 25
            analysis_result["issues"].append("Missing essential sections (Experience, Education, Skills)")
        
        # Check resume length
        if len(words) < 200:
            analysis_result["ats_score"] -= 20
            analysis_result["issues"].append("Resume appears too short (less than 200 words)")
        elif len(words) > 1000:
            analysis_result["ats_score"] -= 10
            analysis_result["issues"].append("Resume might be too long (over 1000 words)")
        
        # Check for contact information
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        phone_pattern = r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b|\b\(\d{3}\)\s?\d{3}[-.]?\d{4}\b'
        
        if not re.search(email_pattern, text_content):
            analysis_result["ats_score"] -= 15
            analysis_result["issues"].append("No email address found")
        
        if not re.search(phone_pattern, text_content):
            analysis_result["ats_score"] -= 10
            analysis_result["issues"].append("No phone number found")
        
        # Check for problematic formatting elements
        problematic_found = 0
        for element in self.problematic_elements:
            if element in text_lower:
                problematic_found += 1
        
        if problematic_found > 0:
            analysis_result["ats_score"] -= min(15, problematic_found * 5)
            analysis_result["issues"].append("Contains elements that may not be ATS-friendly (tables, images, etc.)")
        
        # Check for proper heading structure
        heading_patterns = [
            r'\b(professional\s+)?experience\b',
            r'\beducation\b',
            r'\b(technical\s+)?skills\b',
            r'\b(professional\s+)?summary\b'
        ]
        
        headings_found = 0
        for pattern in heading_patterns:
            if re.search(pattern, text_lower):
                headings_found += 1
        
        if headings_found < 3:
            analysis_result["ats_score"] -= 10
            analysis_result["issues"].append("Missing clear section headings")
        
        # Check for dates (work experience validation)
        date_patterns = [
            r'\b(19|20)\d{2}\b',  # Years
            r'\b(jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\w*\s+(19|20)\d{2}\b',  # Month Year
            r'\b\d{1,2}\/\d{1,2}\/\d{2,4}\b'  # Date format
        ]
        
        dates_found = 0
        for pattern in date_patterns:
            dates_found += len(re.findall(pattern, text_lower))
        
        if dates_found < 2:
            analysis_result["ats_score"] -= 10
            analysis_result["issues"].append("Insufficient date information for work experience")
        
        # Check for action verbs and quantifiable achievements
        action_verbs = [
            'achieved', 'managed', 'led', 'developed', 'created', 'implemented',
            'improved', 'increased', 'reduced', 'organized', 'coordinated'
        ]
        
        action_verb_count = 0
        for verb in action_verbs:
            action_verb_count += len(re.findall(r'\b' + verb + r'\w*\b', text_lower))
        
        if action_verb_count < 3:
            analysis_result["ats_score"] -= 10
            analysis_result["issues"].append("Few action verbs found - use more dynamic language")
        
        # Check for quantifiable achievements (numbers/percentages)
        numbers_pattern = r'\b\d+(%|\+|k|m|million|thousand|dollars?|\$)\b'
        quantifiable_achievements = len(re.findall(numbers_pattern, text_lower))
        
        if quantifiable_achievements < 2:
            analysis_result["ats_score"] -= 8
            analysis_result["issues"].append("Include more quantifiable achievements with specific numbers")
        
        # Check for industry-relevant keywords
        tech_keyword_count = sum(1 for keyword in self.tech_keywords if keyword in text_lower)
        business_keyword_count = sum(1 for keyword in self.business_keywords if keyword in text_lower)
        
        total_keyword_count = tech_keyword_count + business_keyword_count
        if total_keyword_count < 3:
            analysis_result["ats_score"] -= 12
            analysis_result["issues"].append("Consider adding more industry-relevant keywords")
        
        # Check for formatting issues that ATS systems struggle with
        special_chars = len(re.findall(r'[•◦▪▫►]', text_content))  # Bullet points
        if special_chars > 20:
            analysis_result["ats_score"] -= 5
            analysis_result["issues"].append("Too many special characters - use simple bullet points")
        
        # Check for proper spacing and formatting
        lines = text_content.split('\n')
        empty_lines = sum(1 for line in lines if not line.strip())
        if empty_lines > len(lines) * 0.3:
            analysis_result["ats_score"] -= 5
            analysis_result["issues"].append("Excessive white space may cause parsing issues")
        
        # Ensure minimum score
        analysis_result["ats_score"] = max(analysis_result["ats_score"], 0)
        
        # Generate recommendations based on issues
        analysis_result["recommendations"] = self.generate_recommendations(analysis_result["issues"])
        
        return analysis_result

    def generate_recommendations(self, issues):
        """Generate actionable recommendations based on found issues"""
        recommendations = []
        
        for issue in issues:
            if "essential sections" in issue.lower():
                recommendations.append("Add clear sections: Professional Experience, Education, and Skills")
            elif "too short" in issue.lower():
                recommendations.append("Expand your resume with more detailed job descriptions and achievements")
            elif "too long" in issue.lower():
                recommendations.append("Consider condensing information to 1-2 pages maximum")
            elif "email" in issue.lower():
                recommendations.append("Add a professional email address in your contact information")
            elif "phone" in issue.lower():
                recommendations.append("Include your phone number in a standard format")
            elif "ats-friendly" in issue.lower():
                recommendations.append("Avoid tables, images, and complex formatting - use simple text")
            elif "section headings" in issue.lower():
                recommendations.append("Use standard headings like 'Experience', 'Education', 'Skills'")
            elif "date information" in issue.lower():
                recommendations.append("Include clear dates for your work experience and education")
            elif "action verbs" in issue.lower():
                recommendations.append("Start bullet points with strong action verbs like 'managed', 'developed', 'achieved'")
            elif "quantifiable achievements" in issue.lower():
                recommendations.append("Add specific numbers and metrics to demonstrate your impact")
            elif "keywords" in issue.lower():
                recommendations.append("Include relevant industry keywords and skills from job postings")
        
        return recommendations


@require_http_methods(["GET", "POST"])
def resumeats(request):
    """Handle ATS resume checking requests"""
    
    if request.method == "GET":
        return render(request, "resumeats.html")
    
    # Handle POST request (file upload)
    try:
        uploaded_file = request.FILES.get('resume')
        
        if not uploaded_file:
            return JsonResponse({
                "error": "No file uploaded. Please select a resume file."
            }, status=400)
        
        # Validate file extension
        file_ext = os.path.splitext(uploaded_file.name)[1].lower()
        if file_ext not in ALLOWED_EXTENSIONS:
            return JsonResponse({
                "error": f"Unsupported file type. Please upload PDF, DOCX, or TXT files only."
            }, status=400)
        
        # Validate file size
        if uploaded_file.size > MAX_FILE_SIZE:
            return JsonResponse({
                "error": "File too large. Please upload a file smaller than 5MB."
            }, status=400)
        
        # Check if required libraries are available
        if file_ext == ".pdf" and not HAS_PDF_SUPPORT:
            return JsonResponse({
                "error": "PDF processing not available. Please install PyPDF2 or upload a different format."
            }, status=500)
        
        if file_ext == ".docx" and not HAS_DOCX_SUPPORT:
            return JsonResponse({
                "error": "DOCX processing not available. Please install python-docx or upload a different format."
            }, status=500)
        
        # Initialize ATS checker
        checker = ATSResumeChecker()
        
        # Extract text from file
        try:
            text_content = checker.extract_text_from_file(uploaded_file, file_ext)
        except Exception as e:
            return JsonResponse({
                "error": f"Error reading file: {str(e)}"
            }, status=500)
        
        # Analyze the resume
        analysis_result = checker.analyze_resume(text_content, uploaded_file.name)
        
        return JsonResponse(analysis_result, status=200)
        
    except Exception as e:
        return JsonResponse({
            "error": f"An unexpected error occurred: {str(e)}"
        }, status=500)


# Optional: Additional utility functions
def get_file_info(uploaded_file):
    """Get detailed file information"""
    file_info = {
        "name": uploaded_file.name,
        "size": uploaded_file.size,
        "content_type": uploaded_file.content_type,
        "extension": os.path.splitext(uploaded_file.name)[1].lower()
    }
    return file_info


def save_uploaded_file(uploaded_file, subfolder="resumes"):
    """Optionally save uploaded files for analysis logging"""
    try:
        file_path = f"{subfolder}/{uploaded_file.name}"
        saved_path = default_storage.save(file_path, ContentFile(uploaded_file.read()))
        return saved_path
    except Exception as e:
        print(f"Error saving file: {e}")
        return None